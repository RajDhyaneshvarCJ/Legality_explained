import io
import pdfplumber
import docx


class ExtractionError(Exception):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "doc":
        raise ExtractionError(
            "legacy .doc files are not supported - open it in word and save as .docx, then re-upload"
        )
    # assume plain text
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text() for p in pdf.pages]
    except Exception:
        raise ExtractionError("could not read this pdf - it may be corrupted or password protected")

    text = [p for p in pages if p]
    if not text:
        raise ExtractionError(
            "this pdf has no extractable text - it's likely a scanned image. ocr is not supported in v1"
        )
    return "\n".join(text)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception:
        raise ExtractionError("could not read this docx - it may be corrupted")

    parts = [para.text for para in doc.paragraphs if para.text.strip()]

    # contracts often put key numbers in tables so

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise ExtractionError("this docx contains no readable text")
    return "\n".join(parts)
